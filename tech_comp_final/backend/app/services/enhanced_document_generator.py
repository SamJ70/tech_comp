# backend/app/services/enhanced_document_generator.py
"""
ImprovedDocumentGenerator (Full-featured)
- Backwards compatible with generate_document(..., raw_data=...)
- Produces richer DOCX reports with:
    * Title page and metadata
    * Executive summary (existing + auto-generated signals)
    * Per-country detailed sections: risk, compliance, counts, top items
    * Year-wise timeline table (chronological_analysis/chronological_tracking)
    * Wassenaar matches with excerpts
    * News classification (military vs civil) and counts
    * Recommendations & Action Items
    * Charts (source counts, timeline) when matplotlib available
    * Sources & provenance appendix
"""
import os
import io
import logging
from typing import Dict, Any, Optional, List
from datetime import datetime

from app.config import PLOTS_DIR, REPORTS_DIR

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

# Optional dependencies
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except Exception as e:
    Document = None
    DOCX_AVAILABLE = False
    logger.warning("python-docx missing: %s", e)

try:
    import matplotlib.pyplot as plt
    MATPLOTLIB_AVAILABLE = True
except Exception as e:
    plt = None
    MATPLOTLIB_AVAILABLE = False
    logger.info("matplotlib missing: %s", e)

os.makedirs(PLOTS_DIR, exist_ok=True)
os.makedirs(REPORTS_DIR, exist_ok=True)


class ImprovedDocumentGenerator:
    def __init__(self, author: str = "Tech Intelligence Platform"):
        self.author = author

    # ---------------------------
    # Helper / defensive methods
    # ---------------------------
    def _ensure_list(self, maybe):
        if maybe is None:
            return []
        if isinstance(maybe, (list, tuple)):
            return list(maybe)
        return [maybe]

    def _ensure_list_of_dicts(self, maybe):
        out = []
        for itm in self._ensure_list(maybe):
            if isinstance(itm, dict):
                out.append(itm)
            else:
                # convert strings / primitives into dict with text
                out.append({"text": str(itm)})
        return out

    def _safe_get(self, d, k, default=None):
        if not isinstance(d, dict):
            return default
        return d.get(k, default)

    def _short(self, s, n=300):
        if not s:
            return ""
        s = str(s)
        return s if len(s) <= n else s[:n].rstrip() + "..."

    # simple keyword-based news classifier (military vs civil)
    def _classify_news_item(self, title: str, abstract: str):
        text = f"{title or ''} {abstract or ''}".lower()
        military_terms = [
            "military", "army", "navy", "air force", "defence", "defense", "weapon", "missile",
            "torpedo", "drone strike", "unmanned", "combat", "warfare", "militar", "munition",
            "sanction", "ballistic", "armour", "arms", "weaponization", "dual-use"
        ]
        for t in military_terms:
            if t in text:
                return "military"
        # fallback heuristics: words implying research/civil
        civil_terms = ["policy", "research", "study", "commercial", "industry", "education", "climate", "health", "energy"]
        for t in civil_terms:
            if t in text:
                return "civil"
        return "uncertain"

    # chart helpers
    def _create_bar_chart(self, labels: List[str], values: List[int], filename: str, figsize=(5, 2)):
        if not MATPLOTLIB_AVAILABLE:
            logger.debug("matplotlib not available; skipping chart creation")
            return None
        try:
            fig, ax = plt.subplots(figsize=figsize, dpi=120)
            y_pos = range(len(labels))[::-1]
            ax.barh(range(len(labels)), values, align='center')
            ax.set_yticks(range(len(labels)))
            ax.set_yticklabels(labels)
            ax.invert_yaxis()
            ax.set_xlabel("Count")
            plt.tight_layout()
            fullpath = os.path.join(PLOTS_DIR, filename)
            fig.savefig(fullpath, bbox_inches='tight', dpi=150)
            plt.close(fig)
            return fullpath
        except Exception as e:
            logger.exception("chart creation failed: %s", e)
            return None

    def _create_line_chart(self, years: List[int], values: List[int], filename: str, figsize=(6, 2)):
        if not MATPLOTLIB_AVAILABLE:
            return None
        try:
            fig, ax = plt.subplots(figsize=figsize, dpi=120)
            ax.plot(years, values, marker='o')
            ax.set_xlabel("Year")
            ax.set_ylabel("Events")
            ax.grid(axis='y', linestyle='--', alpha=0.4)
            plt.tight_layout()
            fullpath = os.path.join(PLOTS_DIR, filename)
            fig.savefig(fullpath, bbox_inches='tight', dpi=150)
            plt.close(fig)
            return fullpath
        except Exception as e:
            logger.exception("line chart creation failed: %s", e)
            return None

    # ---------------------------
    # Document assembly
    # ---------------------------
    def generate_document(
        self,
        country1: str,
        country2: Optional[str],
        domain: str,
        analysis: Dict[str, Any],
        filepath: str,
        include_charts: bool = True,
        raw_data: Optional[Dict[str, Any]] = None
    ):
        """
        Main entry. Expects:
        - analysis: pipeline output (type comparison or single_country)
        - raw_data: optional raw_data dict keyed by country -> {publications, news, tim, aspi, patents, raw_text}
        """
        raw_data = raw_data or {}
        analysis = analysis or {}

        # Normalize raw_data per country to list-of-dicts form
        for k, v in list(raw_data.items()):
            raw_data[k] = {
                "publications": self._ensure_list_of_dicts(v.get("publications") if isinstance(v, dict) else v and (v if isinstance(v, dict) else {} ).get("publications", []) if isinstance(v, dict) else []),
                "news": self._ensure_list_of_dicts(v.get("news") if isinstance(v, dict) else []),
                "tim": self._ensure_list_of_dicts(v.get("tim") if isinstance(v, dict) else []),
                "aspi": self._ensure_list_of_dicts(v.get("aspi") if isinstance(v, dict) else []),
                "patents": self._ensure_list_of_dicts(v.get("patents") if isinstance(v, dict) else []),
                "raw_text": self._ensure_list_of_dicts(v.get("raw_text") if isinstance(v, dict) else [])
            }

        # fallback for analysis fields
        meta = analysis.get("metadata") or {}
        overall = analysis.get("overall_analysis") or analysis.get("summary") or ""
        nowstr = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # If python-docx not available, create a plain text file as backup but provide rich content
        if not DOCX_AVAILABLE:
            try:
                with open(filepath, "w", encoding="utf-8") as f:
                    f.write(f"Strategic Tech Tracker — {domain}\nGenerated: {nowstr}\n\n")
                    f.write("EXECUTIVE SUMMARY\n")
                    f.write(str(overall) + "\n\n")
                    f.write("ANALYSIS DICT (raw):\n")
                    f.write(str(analysis))
                logger.info("python-docx not installed; wrote text fallback report to %s", filepath)
                return
            except Exception as e:
                logger.exception("Failed to write fallback text report: %s", e)
                return

        # Build DOCX
        doc = Document()
        # Basic styles tweak (optional)
        try:
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
            style.font.size = Pt(11)
        except Exception:
            pass

        # Title page
        doc.add_heading(f"{domain} — {country1}" + (f" vs {country2}" if country2 else " — Single country"), level=0)
        doc.add_paragraph(f"Generated: {nowstr}")
        if meta.get("analyzed_at"):
            doc.add_paragraph(f"Analyzed at: {meta.get('analyzed_at')}")
        doc.add_paragraph("")  # space

        # Executive summary (prefer existing)
        doc.add_heading("Executive summary", level=1)
        exec_text = overall if overall else self._auto_generate_exec_summary(country1, country2, domain, analysis, raw_data)
        doc.add_paragraph(self._short(exec_text, 2500))

        # Quick signals row
        doc.add_heading("Quick signals", level=2)
        # Build a quick table: country, risk, compliance, matched_count, publications, news
        if (analysis.get("type") == "comparison") or country2:
            countries = analysis.get("countries") or [country1] if not country2 else [country1, country2]
        else:
            countries = [country1]
        quick = doc.add_table(rows=1, cols=6)
        quick.rows[0].cells[0].text = "Country"
        quick.rows[0].cells[1].text = "Risk"
        quick.rows[0].cells[2].text = "Compliance"
        quick.rows[0].cells[3].text = "Matched"
        quick.rows[0].cells[4].text = "Publications"
        quick.rows[0].cells[5].text = "News (military/civil/uncertain)"
        for c in countries:
            row = quick.add_row().cells
            row[0].text = str(c)
            c_dual = (analysis.get("dual_use_analysis") or {}).get(c) or (analysis.get("dual_use_analysis") if isinstance(analysis.get("dual_use_analysis"), dict) and not analysis.get("dual_use_analysis").get(c) else {})
            risk = (c_dual or {}).get("risk_level") or "UNKNOWN"
            comp = (c_dual or {}).get("compliance_status") or "UNKNOWN"
            matched = str((c_dual or {}).get("matched_count") or (len((c_dual or {}).get("matched_items") or [])) or 0)
            pubs_count = str(len(raw_data.get(c, {}).get("publications", []) or []))
            # news classification counts
            news_items = raw_data.get(c, {}).get("news", []) or []
            m_count = c_count = u_count = 0
            for n in news_items:
                t = self._safe_get(n, "title", "") or self._safe_get(n, "headline", "") or self._safe_get(n, "text", "")
                a = self._safe_get(n, "abstract", "") or self._safe_get(n, "description", "") or ""
                lab = self._classify_news_item(t, a)
                if lab == "military":
                    m_count += 1
                elif lab == "civil":
                    c_count += 1
                else:
                    u_count += 1
            row[1].text = risk
            row[2].text = comp
            row[3].text = matched
            row[4].text = pubs_count
            row[5].text = f"{m_count}/{c_count}/{u_count}"

        doc.add_paragraph("")  # spacing

        # Per-country details
        for c in countries:
            doc.add_heading(f"{c} — Detailed findings", level=1)
            # Risk summary
            c_dual = (analysis.get("dual_use_analysis") or {}).get(c) or (analysis.get("dual_use_analysis") if isinstance(analysis.get("dual_use_analysis"), dict) and not analysis.get("dual_use_analysis").get(c) else {})
            risk_level = (c_dual or {}).get("risk_level", "UNKNOWN")
            compliance_status = (c_dual or {}).get("compliance_status", "UNKNOWN")
            doc.add_paragraph(f"Risk level: {risk_level}    Compliance: {compliance_status}")

            # Matched Wassenaar items (with short excerpt if available)
            matched = (c_dual or {}).get("matched_items") or (c_dual or {}).get("category_breakdown", {}).get("top_matches", {}).get("general") or []
            if matched:
                doc.add_paragraph("Wassenaar / Dual-use matches (sample):")
                t = doc.add_table(rows=1, cols=3)
                t.rows[0].cells[0].text = "Matched keyword/item"
                t.rows[0].cells[1].text = "Source excerpt"
                t.rows[0].cells[2].text = "Severity/notes"
                for mi in matched[:40]:
                    row = t.add_row().cells
                    if isinstance(mi, dict):
                        row[0].text = str(mi.get("matched_keyword") or mi.get("title") or mi.get("text") or "")
                        row[1].text = self._short(mi.get("excerpt") or mi.get("context") or mi.get("text") or "", 300)
                        row[2].text = str(mi.get("severity") or "")
                    else:
                        row[0].text = str(mi)
                        row[1].text = ""
                        row[2].text = ""
            else:
                doc.add_paragraph("No Wassenaar matches detected by heuristic for this country.")

            # Top publications table (title, year, source, url)
            pubs = raw_data.get(c, {}).get("publications") or []
            if pubs:
                doc.add_paragraph("Top Publications (sample):")
                t = doc.add_table(rows=1, cols=4)
                t.rows[0].cells[0].text = "Title"
                t.rows[0].cells[1].text = "Year"
                t.rows[0].cells[2].text = "Source"
                t.rows[0].cells[3].text = "URL"
                for p in pubs[:25]:
                    row = t.add_row().cells
                    row[0].text = self._short(self._safe_get(p, "title", p.get("text") if isinstance(p, dict) else str(p)), 200)
                    row[1].text = str(self._safe_get(p, "year", ""))
                    row[2].text = str(self._safe_get(p, "source", ""))
                    row[3].text = str(self._safe_get(p, "url", ""))
            else:
                doc.add_paragraph("No publications found for this country+domain (local or remote).")

            # News section with classification and top excerpts
            news_items = raw_data.get(c, {}).get("news") or []
            if news_items:
                doc.add_paragraph("News highlights (sample with classification):")
                t = doc.add_table(rows=1, cols=4)
                t.rows[0].cells[0].text = "Headline"
                t.rows[0].cells[1].text = "Source"
                t.rows[0].cells[2].text = "Year"
                t.rows[0].cells[3].text = "Class (mil/civil/uncertain)"
                for n in news_items[:30]:
                    title = self._safe_get(n, "title", "") or self._safe_get(n, "headline", "") or ""
                    source = self._safe_get(n, "source", "") or self._safe_get(n, "publisher", "")
                    year = self._safe_get(n, "year", "") or (self._safe_get(n, "publishedAt", "")[:4] if self._safe_get(n, "publishedAt") else "")
                    label = self._classify_news_item(title, self._safe_get(n, "abstract", "") or self._safe_get(n, "description", ""))
                    row = t.add_row().cells
                    row[0].text = self._short(title, 200)
                    row[1].text = str(source)
                    row[2].text = str(year)
                    row[3].text = label

            # Chronological timeline / yearwise sample
            chrono = (analysis.get("chronological_analysis") or {}).get("timeline") or (analysis.get("chronological_tracking") or {}).get(c, {}).get("timeline") or (analysis.get("chronological_analysis") or {}).get("timeline", []) or []
            if chrono:
                doc.add_paragraph("Year-wise timeline (sample):")
                t = doc.add_table(rows=1, cols=3)
                t.rows[0].cells[0].text = "Year"
                t.rows[0].cells[1].text = "Events"
                t.rows[0].cells[2].text = "Highlights (short)"
                # chrono might be list of dicts with year/total_events/highlights
                for it in chrono[:30]:
                    year = it.get("year") or it.get("date") or it.get("period") or ""
                    events = str(it.get("total_events") or len(it.get("highlights", []) or []))
                    highlights = ""
                    if it.get("highlights"):
                        if isinstance(it.get("highlights"), list):
                            highlights = self._short(", ".join([self._short(h, 120) for h in it.get("highlights")[:3]]), 300)
                        else:
                            highlights = self._short(str(it.get("highlights")), 300)
                    row = t.add_row().cells
                    row[0].text = str(year)
                    row[1].text = events
                    row[2].text = highlights
                # line chart for timeline if numeric years present
                try:
                    years = []
                    vals = []
                    for it in chrono:
                        try:
                            y = int(it.get("year"))
                        except Exception:
                            continue
                        years.append(y)
                        vals.append(int(it.get("total_events") or len(it.get("highlights") or [])))
                    if years and include_charts:
                        chart_path = self._create_line_chart(years, vals, f"{c}_{domain.replace(' ','_')}_timeline.png")
                        if chart_path:
                            doc.add_paragraph("")
                            doc.add_picture(chart_path, width=Inches(6.0))
                except Exception:
                    logger.exception("timeline chart failed for %s", c)

            # Source counts chart
            counts = {
                "publications": len(raw_data.get(c, {}).get("publications", []) or []),
                "news": len(raw_data.get(c, {}).get("news", []) or []),
                "tim": len(raw_data.get(c, {}).get("tim", []) or []),
                "aspi": len(raw_data.get(c, {}).get("aspi", []) or []),
                "patents": len(raw_data.get(c, {}).get("patents", []) or [])
            }
            doc.add_paragraph("Source counts (sample):")
            t = doc.add_table(rows=1, cols=2)
            t.rows[0].cells[0].text = "Source"
            t.rows[0].cells[1].text = "Count"
            for k, v in counts.items():
                row = t.add_row().cells
                row[0].text = k
                row[1].text = str(v)
            if include_charts:
                chart_path = self._create_bar_chart(list(counts.keys()), list(counts.values()), f"{c}_{domain.replace(' ','_')}_counts.png")
                if chart_path:
                    doc.add_paragraph("")
                    doc.add_picture(chart_path, width=Inches(5.5))

            # Recommendations (prefer analyzer-provided)
            recs = (c_dual or {}).get("recommendations") or (analysis.get("recommendations") or [])
            if not recs:
                # auto generate a few standard recommendations
                recs = [
                    "Continue monitoring high-match areas and review raw matched documents manually.",
                    "If matched items indicate potential non-compliance escalate to the oversight team.",
                    "Cross-check patent filings and procurement notices for dual-use signs.",
                    "Prioritize verification of top 5 matched items for intent."
                ]
            doc.add_paragraph("Recommendations and Action Items:")
            for r in recs:
                p = doc.add_paragraph(style=None)
                p.add_run("• ").bold = True
                p.add_run(self._short(r, 400))

            doc.add_page_break()

        # Comparison summary page (if two countries)
        if country2:
            doc.add_heading("Comparison Summary & Tactical Insights", level=1)
            # side-by-side table for key signals
            t = doc.add_table(rows=1, cols=6)
            t.rows[0].cells[0].text = "Signal"
            t.rows[0].cells[1].text = country1
            t.rows[0].cells[2].text = ""
            t.rows[0].cells[3].text = country2
            t.rows[0].cells[4].text = ""
            t.rows[0].cells[5].text = "Notes"

            def cell_for(c, key):
                c_dual = (analysis.get("dual_use_analysis") or {}).get(c) or {}
                if key == "risk":
                    return str(c_dual.get("risk_level") or "UNKNOWN")
                if key == "compliance":
                    return str(c_dual.get("compliance_status") or "UNKNOWN")
                if key == "matched":
                    return str((c_dual.get("matched_count") or len(c_dual.get("matched_items") or [])) if isinstance(c_dual, dict) else "0")
                return ""

            row = t.add_row().cells
            row[0].text = "Risk level"
            row[1].text = cell_for(country1, "risk")
            row[2].text = ""
            row[3].text = cell_for(country2, "risk")
            row[4].text = ""
            row[5].text = ""
            row = t.add_row().cells
            row[0].text = "Compliance"
            row[1].text = cell_for(country1, "compliance")
            row[2].text = ""
            row[3].text = cell_for(country2, "compliance")
            row[4].text = ""
            row[5].text = ""
            row = t.add_row().cells
            row[0].text = "Matched items"
            row[1].text = cell_for(country1, "matched")
            row[2].text = ""
            row[3].text = cell_for(country2, "matched")
            row[4].text = ""
            row[5].text = ""
            doc.add_page_break()

        # Sources & Methodology appendix
        doc.add_heading("Appendix: Sources & Methodology", level=1)
        doc.add_paragraph("This report combined multiple sources (publications, patents, TIM export, ASPI export, news) and user-provided extra sources if included. Wassenaar matches are heuristics based on the provided Wassenaar index and keyword matching; manual review is recommended for any flagged items.")
        doc.add_paragraph("Methodology highlights:")
        doc.add_paragraph("• Data collection: CrossRef, EuropePMC, optional TIM/ASPI local exports, NewsAPI (if configured), local publications/news JSON fallbacks.")
        doc.add_paragraph("• Dual-use matching: heuristic keyword matching to Wassenaar index entries; risk levels derived from matched counts.")
        doc.add_paragraph("• News classification: small keyword classifier to categorize as 'military' vs 'civil' for signal counts (high recall; manual review required).")
        doc.add_paragraph("• Chronology: events aggregated per year from publications/news/TIM/ASPI when available.")

        # Provenance table
        doc.add_heading("Provenance & Data counts", level=2)
        meta = analysis.get("metadata") or {}
        if isinstance(meta.get("sources_used"), dict):
            t = doc.add_table(rows=1, cols=3)
            t.rows[0].cells[0].text = "Country"
            t.rows[0].cells[1].text = "Raw items counted"
            t.rows[0].cells[2].text = "User extra sources"
            for c, cnt in meta.get("sources_used", {}).items():
                row = t.add_row().cells
                row[0].text = str(c)
                row[1].text = str(cnt)
                extra = meta.get("extra_sources_used") or analysis.get("extra_sources_used") or []
                row[2].text = ", ".join(self._ensure_list(extra)) if extra else ""
        else:
            t = doc.add_table(rows=1, cols=2)
            t.rows[0].cells[0].text = "Key"
            t.rows[0].cells[1].text = "Value"
            for k, v in meta.items():
                row = t.add_row().cells
                row[0].text = str(k)
                row[1].text = str(v)

        # Final save
        try:
            doc.save(filepath)
            logger.info("Report saved to %s", filepath)
        except Exception as e:
            logger.exception("Failed to save report: %s", e)

    # ---------------------------
    # Auto-generate a short executive summary if none present
    # ---------------------------
    def _auto_generate_exec_summary(self, country1, country2, domain, analysis, raw_data):
        lines = []
        lines.append(f"Domain: {domain}")
        counts_summary = []
        for c in ([country1] + ([country2] if country2 else [])):
            pubs = len(raw_data.get(c, {}).get("publications", []) or [])
            news = len(raw_data.get(c, {}).get("news", []) or [])
            tim = len(raw_data.get(c, {}).get("tim", []) or [])
            aspi = len(raw_data.get(c, {}).get("aspi", []) or [])
            lines.append(f"{c}: {pubs} publications, {news} news items, {tim} TIM items, {aspi} ASPI items")
            c_dual = (analysis.get("dual_use_analysis") or {}).get(c) or {}
            risk = c_dual.get("risk_level") or "UNKNOWN"
            lines.append(f"  Risk: {risk} (matched items: {(c_dual.get('matched_count') or len(c_dual.get('matched_items') or []))})")
        # top-level suggestion
        lines.append("Topline recommendation: review highest-matched Wassenaar items and prioritize verification of the top 5.")
        return "\n".join(lines)
