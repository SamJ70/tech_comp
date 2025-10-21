from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict
from datetime import datetime

class ImprovedDocumentGenerator:
    def __init__(self):
        pass
    
    def generate_document(
        self, 
        country1: str, 
        country2: str, 
        domain: str, 
        analysis: Dict, 
        filepath: str
    ):
        """Generate comprehensive, evidence-based report"""
        doc = Document()
        
        # Title page
        self._add_title(doc, f"Technology Comparison Report")
        self._add_subtitle(doc, f"{country1} vs {country2}")
        self._add_subtitle(doc, f"Domain: {domain}")
        self._add_metadata(doc)
        
        # Data quality warning if needed
        quality = analysis.get("data_quality", {})
        if quality.get("confidence") != "high":
            self._add_quality_notice(doc, quality)
        
        doc.add_page_break()
        
        # Executive Summary
        self._add_section_header(doc, "Executive Summary")
        doc.add_paragraph(analysis.get("overall_analysis", "Analysis not available."))
        
        doc.add_page_break()
        
        # Concrete Metrics Comparison
        self._add_section_header(doc, "Concrete Metrics Comparison")
        self._add_concrete_metrics_table(doc, country1, country2, analysis)
        
        doc.add_paragraph()
        
        # Detailed comparison
        self._add_section_header(doc, "Detailed Category Comparison")
        comparison = analysis.get("comparison", {})
        if comparison:
            for category, result in comparison.items():
                self._add_subsection_header(doc, category.replace("_", " ").title())
                doc.add_paragraph(result)
        
        doc.add_page_break()
        
        # Country 1 Analysis
        self._add_section_header(doc, f"{country1} - Detailed Profile")
        self._add_country_profile(doc, country1, analysis)
        
        doc.add_page_break()
        
        # Country 2 Analysis
        self._add_section_header(doc, f"{country2} - Detailed Profile")
        self._add_country_profile(doc, country2, analysis)
        
        doc.add_page_break()
        
        # Recent Developments
        self._add_section_header(doc, "Recent Developments & News")
        news_items = analysis.get("news", [])
        if news_items:
            # Group by recency
            recent = [n for n in news_items if n.get("recent", False)]
            older = [n for n in news_items if not n.get("recent", False)]
            
            if recent:
                self._add_subsection_header(doc, "Recent Updates (with dates)")
                for item in recent[:8]:
                    self._add_news_item(doc, item)
            
            if older:
                self._add_subsection_header(doc, "Additional Context")
                for item in older[:5]:
                    self._add_news_item(doc, item)
        else:
            doc.add_paragraph("No recent news items found in sources.")
        
        doc.add_page_break()
        
        # Methodology & Quality
        self._add_section_header(doc, "Methodology & Data Quality")
        self._add_methodology_section(doc, country1, country2, domain, quality)
        
        doc.save(filepath)
    
    def _add_quality_notice(self, doc: Document, quality: Dict):
        """Add prominent data quality notice"""
        warnings = quality.get("warnings", [])
        confidence = quality.get("confidence", "medium")
        
        notice = doc.add_paragraph()
        notice_run = notice.add_run("⚠️ Data Quality Notice\n")
        notice_run.font.bold = True
        notice_run.font.size = Pt(12)
        notice_run.font.color.rgb = RGBColor(204, 102, 0)
        
        conf_text = notice.add_run(f"Confidence Level: {confidence.upper()}\n\n")
        conf_text.font.bold = True
        
        for warning in warnings:
            notice.add_run(f"• {warning}\n")
        
        notice.add_run("\nPlease interpret results with appropriate context.")
        doc.add_paragraph()
    
    def _add_concrete_metrics_table(self, doc: Document, country1: str, country2: str, analysis: Dict):
        """Add table of concrete, verifiable metrics"""
        metrics = analysis.get("concrete_metrics", {})
        metrics1 = metrics.get(country1, {})
        metrics2 = metrics.get(country2, {})
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Headers
        header_cells = table.rows[0].cells
        header_cells[0].text = "Metric"
        header_cells[1].text = country1
        header_cells[2].text = country2
        
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                paragraph.runs[0].font.bold = True
        
        # Funding
        self._add_metric_row(
            table, 
            "Maximum Funding Documented",
            self._format_list(metrics1.get("funding_amounts", [])[:3]),
            self._format_list(metrics2.get("funding_amounts", [])[:3])
        )
        
        # Patents
        self._add_metric_row(
            table,
            "Patent Counts Mentioned",
            self._format_list(metrics1.get("patent_counts", [])[:3]),
            self._format_list(metrics2.get("patent_counts", [])[:3])
        )
        
        # Market Size
        self._add_metric_row(
            table,
            "Market Size References",
            self._format_list(metrics1.get("market_size", [])[:2]),
            self._format_list(metrics2.get("market_size", [])[:2])
        )
        
        # Growth Rates
        self._add_metric_row(
            table,
            "Growth Rates Cited",
            self._format_list(metrics1.get("growth_rates", [])[:3]),
            self._format_list(metrics2.get("growth_rates", [])[:3])
        )
        
        # Research Output
        self._add_metric_row(
            table,
            "Research Papers Count",
            self._format_list(metrics1.get("research_output", [])[:2]),
            self._format_list(metrics2.get("research_output", [])[:2])
        )
        
        # Companies
        companies1 = analysis.get("resources", {}).get(country1, [])
        companies2 = analysis.get("resources", {}).get(country2, [])
        self._add_metric_row(
            table,
            "Organizations Identified",
            f"{len(companies1)} entities",
            f"{len(companies2)} entities"
        )
    
    def _add_metric_row(self, table, metric_name: str, value1: str, value2: str):
        """Add a row to metrics table"""
        row = table.add_row()
        row.cells[0].text = metric_name
        row.cells[1].text = value1 if value1 else "Not found"
        row.cells[2].text = value2 if value2 else "Not found"
    
    def _format_list(self, items: list) -> str:
        """Format list for table cell"""
        if not items:
            return "Not documented"
        return ", ".join(str(item) for item in items[:3])
    
    def _add_country_profile(self, doc: Document, country: str, analysis: Dict):
        """Add detailed country profile"""
        # Summary
        summary = analysis.get("summary", {}).get(country, "No data available.")
        doc.add_paragraph(summary)
        
        doc.add_paragraph()
        
        # Key Organizations
        self._add_subsection_header(doc, "Key Organizations & Entities")
        resources = analysis.get("resources", {}).get(country, [])
        if resources:
            for resource in resources[:10]:
                doc.add_paragraph(resource, style='List Bullet')
        else:
            doc.add_paragraph("No specific organizations identified in available sources.")
        
        doc.add_paragraph()
        
        # Concrete Metrics
        self._add_subsection_header(doc, "Documented Metrics")
        metrics = analysis.get("concrete_metrics", {}).get(country, {})
        
        if metrics.get("funding_amounts"):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run("Funding: ").bold = True
            p.add_run(", ".join(metrics["funding_amounts"][:5]))
        
        if metrics.get("patent_counts"):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run("Patents: ").bold = True
            p.add_run(", ".join(metrics["patent_counts"][:5]))
        
        if metrics.get("market_size"):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run("Market Size: ").bold = True
            p.add_run(", ".join(metrics["market_size"][:3]))
        
        if metrics.get("growth_rates"):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run("Growth Rates: ").bold = True
            p.add_run(", ".join(metrics["growth_rates"][:3]))
        
        if not any([metrics.get("funding_amounts"), metrics.get("patent_counts"), 
                   metrics.get("market_size"), metrics.get("growth_rates")]):
            doc.add_paragraph("Limited quantitative metrics found in available sources.")
    
    def _add_news_item(self, doc: Document, item: Dict):
        """Add a news/highlight item"""
        p = doc.add_paragraph(style='List Bullet')
        
        source = item.get("source", "Source Unknown")
        headline = item.get("headline", "No details available")
        
        source_run = p.add_run(f"[{source}] ")
        source_run.font.bold = True
        source_run.font.color.rgb = RGBColor(0, 102, 204)
        
        p.add_run(headline)
    
    def _add_methodology_section(self, doc: Document, country1: str, country2: str, domain: str, quality: Dict):
        """Add detailed methodology explanation"""
        
        methodology_text = f"""This report compares {country1} and {country2} in the {domain} domain using autonomous data collection and analysis.

**Data Collection Process:**
• Wikipedia search API used to find relevant articles
• Multiple search queries per country ({country1}, {country2}) and domain ({domain})
• Relevance scoring applied to filter content (threshold: 2.0/10.0)
• Sources collected: {quality.get('sources', {}).get(country1, 'N/A')} for {country1}, {quality.get('sources', {}).get(country2, 'N/A')} for {country2}
• Average relevance: {quality.get('relevance_scores', {}).get(country1, 'N/A')}/10 for {country1}, {quality.get('relevance_scores', {}).get(country2, 'N/A')}/10 for {country2}

**Analysis Approach:**
• Concrete metrics extraction: funding amounts, patent counts, market sizes, growth rates
• Company identification with context (major tech companies vs local entities)
• Temporal analysis focusing on developments from 2020-{datetime.now().year}
• Evidence-based comparison using verifiable data points
• Multi-factor scoring system for balanced conclusions

**Key Metrics Extracted:**
• Financial: Funding amounts, market valuations, investment deals
• Innovation: Patent counts, research output, breakthrough mentions
• Ecosystem: Company counts, university presence, government initiatives
• Recent Activity: Dated developments, announcements, launches

**Important Limitations:**
• Data limited to publicly available Wikipedia content in English
• Wikipedia coverage varies by country and topic
• Some countries have more comprehensive documentation than others
• Analysis represents documented information, not comprehensive capabilities
• Classified or proprietary information not included
• Language barriers may affect completeness for non-English-speaking countries

**Confidence Assessment:**
• Overall confidence: {quality.get('confidence', 'medium').upper()}
• Data quality warnings: {len(quality.get('warnings', []))}
"""
        
        if quality.get('warnings'):
            methodology_text += "\n**Specific Data Quality Notes:**\n"
            for warning in quality.get('warnings', []):
                methodology_text += f"• {warning}\n"
        
        methodology_text += f"\n\n**Report Generated:** {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
        methodology_text += "\n**Data Sources:** Wikipedia (via search API and direct access)"
        methodology_text += "\n**Analysis Method:** Automated text analysis with manual validation rules"
        
        doc.add_paragraph(methodology_text)
    
    def _add_title(self, doc: Document, text: str):
        """Add document title"""
        title = doc.add_heading(text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(26)
            run.font.color.rgb = RGBColor(0, 51, 102)
            run.font.bold = True
    
    def _add_subtitle(self, doc: Document, text: str):
        """Add subtitle"""
        subtitle = doc.add_paragraph(text)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in subtitle.runs:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(64, 64, 64)
    
    def _add_metadata(self, doc: Document):
        """Add generation metadata"""
        metadata = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in metadata.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.font.italic = True
    
    def _add_section_header(self, doc: Document, text: str):
        """Add section header"""
        heading = doc.add_heading(text, level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
            run.font.size = Pt(18)
    
    def _add_subsection_header(self, doc: Document, text: str):
        """Add subsection header"""
        heading = doc.add_heading(text, level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(51, 102, 153)
            run.font.size = Pt(14)